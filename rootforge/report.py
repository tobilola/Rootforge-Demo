"""Generate the human-review investigation report.

The report carries its own provenance: which model produced it, against which
documents, and how many claims passed grounding verification. An investigation
record that cannot say where it came from is not an investigation record.
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt

STATUS_MARK = {
    "verified": "",
    "paraphrase": "",
    "unverified": "  [UNVERIFIED — REVIEWER MUST CONFIRM]",
    "bad_ref": "  [CITATION ERROR — REVIEWER MUST CONFIRM]",
}


def _cite(item):
    if not isinstance(item, dict):
        return " [Source: no citation]"
    refs = "; ".join(item.get("refs", [])) or "no citation"
    verification = item.get("verification", {})
    if not isinstance(verification, dict):
        verification = {}
    mark = STATUS_MARK.get(verification.get("status", ""), "")
    return f" [Source: {refs}]{mark}"


def build_report(case, investigation, documents, reviewer, decision, notes):
    doc = Document()
    doc.add_heading("Rootforge Investigation Draft", 0)
    p = doc.add_paragraph(
        "Synthetic demonstration. AI-assisted draft for qualified human review. "
        "Not a final quality decision and not a determination of root cause."
    )
    p.runs[0].font.size = Pt(9)

    doc.add_heading(f"{case['id']}: {case['title']}", level=1)

    run = investigation.get("run", {})
    ver = investigation.get("verification", {})
    doc.add_heading("Provenance", level=2)
    for line in [
        f"Model: {run.get('model', 'unknown')} ({run.get('provider', 'unknown')})",
        f"Documents analysed: {len(documents)} — {', '.join(d['ref'] for d in documents)}",
        f"Generation time: {run.get('seconds', '?')} seconds",
        f"Grounding verification: {ver.get('summary', 'not run')}",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    failures = ver.get("failures", [])
    if failures:
        doc.add_heading("Flagged model assertions — excluded from confirmed findings", level=1)
        doc.add_paragraph(
            "These statements did not resolve to source text. They are excluded "
            "from confirmed findings and retained for reviewer inspection."
        )
        for f in failures:
            doc.add_paragraph(
                f"{f.get('text', '')} — {f.get('detail', 'Grounding check failed')}",
                style="List Bullet",
            )

    doc.add_heading("Evidence-linked timeline", level=1)
    for item in investigation.get("timeline", []):
        doc.add_paragraph(
            f"{item.get('time','')} — {item.get('text','')}"
            f" (confidence: {item.get('confidence','?')})" + _cite(item),
            style="List Bullet",
        )

    for title, key in [
        ("Confirmed by the record", "confirmed"),
        ("Missing from the record", "missing"),
        ("Contradictions and doubts", "contradictions"),
    ]:
        doc.add_heading(title, level=1)
        items = investigation.get("findings", {}).get(key, [])
        if key == "confirmed":
            items = [item for item in items if isinstance(item, dict) and
                     item.get("verification", {}).get("status")
                     not in ("unverified", "bad_ref")]
        if not items:
            doc.add_paragraph("None identified.")
        for item in items:
            doc.add_paragraph(item.get("text", "") + _cite(item), style="List Bullet")

    doc.add_heading("Root-cause hypotheses (not conclusions)", level=1)
    doc.add_paragraph(
        "The following are proposed lines of enquiry. None has been established. "
        "Each states what evidence would disconfirm it."
    )
    for i, h in enumerate(investigation.get("hypotheses", []), 1):
        doc.add_paragraph(
            f"H{i}: {h.get('statement','')} "
            f"({h.get('confidence','?')} confidence; {h.get('status','')})" + _cite(h),
            style="List Bullet",
        )
        doc.add_paragraph(f"    Reasoning: {h.get('reasoning','')}")
        doc.add_paragraph(f"    Would be disconfirmed by: {h.get('disconfirming_evidence','')}")

    doc.add_heading("Draft corrective and preventive actions", level=1)
    for a in investigation.get("capa", []):
        if not isinstance(a, dict):
            continue
        owner = a.get("owner_role") or a.get("owner", "Unassigned")
        due = f" Due: {a['due']}." if a.get("due") else ""
        doc.add_paragraph(
            f"{a.get('type','Action')}: {a.get('text') or a.get('action','')} "
            f"Owner: {owner}.{due}" + _cite(a),
            style="List Bullet",
        )
        if a.get("rationale"):
            doc.add_paragraph(f"    Rationale: {a['rationale']}")

    doc.add_heading("Human review", level=1)
    doc.add_paragraph(f"Reviewer: {reviewer or 'Not entered'}")
    doc.add_paragraph(f"Decision: {decision}")
    doc.add_paragraph(f"Notes: {notes or 'None'}")
    doc.add_paragraph(
        "All hypotheses and actions above remain subject to qualified human "
        "approval. This document does not constitute a regulatory determination."
    )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
