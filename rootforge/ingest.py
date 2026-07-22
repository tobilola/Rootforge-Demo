"""Turn uploaded evidence into the corpus format the engine expects.

Supports pdf, docx, txt and md. Each uploaded file becomes one citable
document whose `ref` is derived from the filename, so citations in the output
point at something the customer recognises.
"""

from __future__ import annotations

import io
import re


def _ref_from_name(name: str, taken: set[str]) -> str:
    stem = re.sub(r"\.[^.]+$", "", name)
    ref = re.sub(r"[^A-Za-z0-9]+", "-", stem).strip("-").upper()[:24] or "DOC"
    base, n = ref, 2
    while ref in taken:
        ref, n = f"{base}-{n}", n + 1
    taken.add(ref)
    return ref


def _read_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _read_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(p for p in parts if p.strip())


def load_uploads(files) -> list[dict]:
    """`files` is a list of Streamlit UploadedFile (or any .name/.read() pair)."""
    docs, taken = [], set()
    for f in files:
        data = f.read()
        name = f.name
        low = name.lower()
        try:
            if low.endswith(".pdf"):
                text = _read_pdf(data)
            elif low.endswith(".docx"):
                text = _read_docx(data)
            else:
                text = data.decode("utf-8", errors="replace")
        except Exception as exc:  # surface, don't swallow
            text = ""
            name = f"{name} [unreadable: {exc}]"
        if not text.strip():
            continue
        docs.append({
            "ref": _ref_from_name(f.name, taken),
            "name": f.name,
            "kind": "uploaded evidence",
            "text": text.strip(),
        })
    return docs


def corpus_from(documents: list[dict]) -> str:
    return "\n\n".join(
        f"<document ref=\"{d['ref']}\" name=\"{d['name']}\" kind=\"{d['kind']}\">\n"
        f"{d['text']}\n</document>"
        for d in documents
    )
